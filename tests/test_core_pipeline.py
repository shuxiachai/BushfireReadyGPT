import json
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from textwrap import dedent
from types import SimpleNamespace
from zipfile import ZipFile

import httpx
import pytest
from docx import Document
from openai import APIConnectionError, APIStatusError
from pypdf import PdfReader

from src import export_package as export_package_module
from src import session_store
from src.agents import run_analysis_pipeline
from src.agents.community_vulnerability_agent import CommunityVulnerabilityAgent
from src.agents.profile_agent import ProfileAgent
from src.agents.report_quality_agent import ReportQualityAgent
from src.audit import save_report_audit
from src.coverage_map import get_coverage_table, load_coverage_geojson
from src.data_paths import PROJECT_ROOT, DataPaths
from src.data_status import get_community_data_status
from src.docx_export import create_report_docx
from src.export_content import extract_report_metadata
from src.export_package import PILOT_EXPORT_SCHEMA, create_pilot_export_package
from src.export_register import build_export_register_snapshot
from src.governance import DRAFT_STATUS, build_review_checklist_snapshot
from src.licence_register import get_licence_register, licence_register_csv
from src.model_runtime import model_service_error_message
from src.pdf_export import create_report_pdf
from src.report_generation_quality import (
    QUALITY_POLICY_FINGERPRINT,
    QUALITY_POLICY_VERSION,
    assess_generated_narrative,
    attributed_rag_source_ids,
    build_report_repair_prompt,
    evaluate_governed_report,
    normalize_generated_narrative,
    structural_gate_passed,
)
from src.report_template import (
    append_evidence_tables,
    append_human_signoff,
    apply_governance_notice,
    build_human_signoff,
)
from src.report_workflow import (
    validate_geography_consistency,
    validate_report_inputs,
    validate_review_record,
)
from src.safety_boundary import evaluate_safety_boundaries
from src.source_attribution import format_official_attribution, format_rag_attribution

SAMPLE_REPORT = """# Cairns Campus Preparedness Report

Location: Cairns, Queensland
Audience: students and teachers

## Executive Summary
Preparedness planning draft.

| Field | Value |
| --- | --- |
| Scenario | School preparedness |

## Evidence Confidence and Provenance

| Code | Evidence class | Confidence / use boundary | Required review |
| --- | --- | --- | --- |
| O1 | Official-source reference | Current information not confirmed | Check the official source |
| P2 | Processed official-origin data | Transformation limitations apply | Check source year and method |
| R3 | Deterministic rule inference | Indicative planning logic | Validate locally |
| A4 | AI-generated draft synthesis | Not an evidence source | Human verification required |
| U0 | User-provided context | Unverified | Confirm organisational records |

- [ ] Confirm official sources.
"""


def _allow_governed_package_export(monkeypatch):
    def passing_quality(_report, _analysis):
        return {
            "checks": [],
            "summary": {"passed": 1, "warnings": 0, "failed": 0, "total": 1},
            "approval_gate": {"passed": True, "status": "passed", "blocking_failures": []},
            "assessment_scope": "Test fixture",
            "quality_policy_version": QUALITY_POLICY_VERSION,
            "quality_policy_fingerprint": QUALITY_POLICY_FINGERPRINT,
        }

    monkeypatch.setattr("src.audit.evaluate_governed_report", passing_quality)
    monkeypatch.setattr(export_package_module, "evaluate_governed_report", passing_quality)


def test_analysis_pipeline_returns_expected_sections():
    analysis = run_analysis_pipeline(
        location="Cairns, Queensland",
        audience="students and teachers",
        scenario="School Preparedness",
        concerns=["evacuation", "assembly points", "first aid"],
        timeframe="One-week action plan",
        extra_context="Campus emergency planning pilot.",
    )

    expected_keys = {
        "profile",
        "data",
        "community",
        "risk_context",
        "plan",
        "prompt_context",
        "evidence_confidence",
    }
    assert expected_keys.issubset(analysis)
    assert analysis["profile"]["location"] == "Cairns, Queensland"
    assert analysis["prompt_context"]
    assert {row["code"] for row in analysis["evidence_confidence"]} == {
        "O1",
        "P2",
        "R3",
        "A4",
        "U0",
    }


def test_report_appendices_are_idempotent():
    analysis = run_analysis_pipeline(
        location="Cairns, Queensland",
        audience="students and teachers",
        scenario="School Preparedness",
        concerns=["official sources"],
        timeframe="One-week action plan",
        extra_context="",
    )

    report = apply_governance_notice("# Test Report")
    report = append_evidence_tables(report, analysis)
    report = append_evidence_tables(report, analysis)
    report = append_human_signoff(report, {"reviewer_name": "Test Reviewer"})

    assert report.count("DRAFT STATUS NOTICE") == 1
    assert report.count("## Evidence Tables") == 1
    assert report.count("### Evidence Confidence and Provenance") == 1
    assert "[O1]" in report
    assert "[P2]" in report
    assert "[R3]" in report
    assert "[A4]" in report
    assert "[U0]" in report
    assert "## Human Review Sign-off" in report
    assert "Test Reviewer" in report


def test_deterministic_evidence_tables_replace_model_modified_tables():
    analysis = run_analysis_pipeline(
        location="Cairns, Queensland",
        audience="students and teachers",
        scenario="School Preparedness",
        concerns=["official sources"],
        timeframe="One-week action plan",
        extra_context="",
    )
    model_modified = """# Draft

## Evidence Tables

MODEL MODIFIED EVIDENCE THAT MUST NOT SURVIVE

## Human Review Sign-off

Old sign-off.
"""

    report = append_evidence_tables(model_modified, analysis)
    report = append_human_signoff(report, {"reviewer_name": "Test Reviewer"})

    assert "MODEL MODIFIED EVIDENCE" not in report
    assert report.count("## Evidence Tables") == 1
    assert report.index("## Evidence Tables") < report.index("## Human Review Sign-off")


def test_governance_notice_replaces_modified_notice_without_losing_report_body():
    modified = """**DRAFT STATUS NOTICE**

Modified model wording that must not survive.

# Preparedness Report

## Executive Summary
Keep this report body.
"""

    report = apply_governance_notice(modified)

    assert report.count("DRAFT STATUS NOTICE") == 1
    assert "Modified model wording" not in report
    assert "# Preparedness Report" in report
    assert "Keep this report body" in report


def test_session_state_persistence_uses_json_and_round_trips(monkeypatch):
    with TemporaryDirectory() as directory:
        target = Path(directory) / "session.json"
        state = {
            "messages": [{"role": "assistant", "content": "Draft", "kind": "report"}],
            "latest_report": {"id": "report-1", "version": 2, "text": "Draft"},
        }
        monkeypatch.setattr(session_store, "SESSION_STATE_PATH", str(target))
        monkeypatch.setattr(session_store, "st", SimpleNamespace(session_state=state))

        session_store.persist_session_state()
        saved = json.loads(target.read_text(encoding="utf-8"))

        assert saved["latest_report"]["version"] == 2
        assert session_store._load_persisted_state() == saved


def test_data_and_licence_registers_load():
    status = get_community_data_status()
    licences = get_licence_register()
    licence_csv = licence_register_csv()

    assert status["active_exists"] is True
    assert status["row_count"] >= 1
    assert status["source_period"] == "2021 Census and 2022 ERP fields"
    assert status["latest_source_year"] == 2022
    assert status["source_age_years"] >= 0
    assert licences["licence_register"]
    assert "source_name" in licence_csv


def test_default_data_paths_do_not_depend_on_the_working_directory(monkeypatch, tmp_path):
    for environment_name in (
        "BUSHFIRE_DATA_DIR",
        "BUSHFIRE_COMMUNITY_PROFILE_PATH",
        "BUSHFIRE_OFFICIAL_SOURCES_PATH",
        "BUSHFIRE_REGION_MAPPINGS_PATH",
        "BUSHFIRE_RISK_CONTEXT_RULES_PATH",
    ):
        monkeypatch.delenv(environment_name, raising=False)
    monkeypatch.chdir(tmp_path)

    paths = DataPaths.from_env()
    status = get_community_data_status()
    analysis = run_analysis_pipeline(
        location="Cairns, Queensland",
        audience="community residents",
        scenario="Community preparedness",
        concerns=["Evacuation"],
        timeframe="7-day action plan",
        extra_context="Working-directory regression test.",
    )

    assert paths.data_dir == (PROJECT_ROOT / "data_australia").resolve()
    assert Path(status["active_path"]).is_relative_to(PROJECT_ROOT)
    assert analysis["community"]["matched_location"] == "Cairns, Queensland"
    assert get_coverage_table(location_filter="Cairns")


def test_runtime_community_path_override_is_shared_by_map_agent_and_status(
    monkeypatch,
    tmp_path,
):
    data_dir = tmp_path / "configured-data"
    override_path = tmp_path / "overrides" / "community.csv"
    data_dir.mkdir(parents=True)
    (data_dir / "official_sources.yml").write_text("sources: []\n", encoding="utf-8")
    (data_dir / "region_mappings.yml").write_text(
        "regions:\n  - location: Testville\n    state: New South Wales\n",
        encoding="utf-8",
    )
    (data_dir / "risk_context_rules.yml").write_text("rules: []\n", encoding="utf-8")
    override_path.parent.mkdir(parents=True)
    override_path.write_text(
        "location,state,population,older_people_pct,no_car_households_pct,"
        "language_support_needed,risk_notes,source,source_years\n"
        "Testville,New South Wales,1234,18,9,medium,Local test risk,Test source,2026\n",
        encoding="utf-8",
    )
    monkeypatch.setenv("BUSHFIRE_DATA_DIR", str(data_dir))
    monkeypatch.setenv("BUSHFIRE_COMMUNITY_PROFILE_PATH", str(override_path))

    paths = DataPaths.from_env()
    map_rows = get_coverage_table()
    agent_result = CommunityVulnerabilityAgent().run(
        {
            "location": "Testville, New South Wales",
            "state": "New South Wales",
        }
    )
    pipeline_result = run_analysis_pipeline(
        location="Testville, New South Wales",
        audience="community residents",
        scenario="Community preparedness",
        concerns=["Evacuation"],
        timeframe="7-day action plan",
        extra_context="Injected path regression test.",
        data_paths=paths,
    )
    status = get_community_data_status()

    assert paths.data_dir == data_dir.resolve()
    assert paths.risk_context_rules == (data_dir / "risk_context_rules.yml").resolve()
    assert paths.community_profile == override_path.resolve()
    assert map_rows[0]["location"] == "Testville"
    assert agent_result["matched_location"] == "Testville, New South Wales"
    assert agent_result["data_quality"]["source_period"] == "2026"
    assert pipeline_result["community"]["matched_location"] == "Testville, New South Wales"
    assert pipeline_result["community"]["data_quality"]["latest_source_year"] == 2026
    assert status["active_path"] == str(override_path.resolve())
    assert status["locations"] == ["Testville"]
    assert status["freshness"] == "Unknown source age"


def test_coverage_geojson_cache_is_scoped_to_the_resolved_path(monkeypatch, tmp_path):
    first_path = tmp_path / "first.geojson"
    second_path = tmp_path / "second.geojson"
    first_path.write_text(
        json.dumps({"type": "FeatureCollection", "source": "first", "features": []}),
        encoding="utf-8",
    )
    second_path.write_text(
        json.dumps({"type": "FeatureCollection", "source": "second", "features": []}),
        encoding="utf-8",
    )

    monkeypatch.setenv("BUSHFIRE_SA2_COVERAGE_PATH", str(first_path))
    first = load_coverage_geojson()
    monkeypatch.setenv("BUSHFIRE_SA2_COVERAGE_PATH", str(second_path))
    second = load_coverage_geojson()

    assert first["source"] == "first"
    assert second["source"] == "second"


def test_pilot_export_package_contains_governance_files(tmp_path, monkeypatch):
    _allow_governed_package_export(monkeypatch)
    review_record = {
        "approval_status": DRAFT_STATUS,
        "reviewer_name": "Test Reviewer",
    }
    report_text = append_human_signoff("# Test Report\n\nPreparedness draft.", review_record)
    package_context = {
        "location": "Cairns",
        "report_status": "Draft - human review required",
        "report_id": "governance-files-test",
        "report_version": 1,
    }
    register_snapshot = build_export_register_snapshot()
    monkeypatch.setenv("BUSHFIRE_AUDIT_DIR", str(tmp_path))
    audit_path = save_report_audit(
        {
            "report_id": "governance-files-test",
            "report_version": 1,
            "report_text": report_text,
            "inputs": {"location": "Cairns"},
            "human_review": review_record,
            "export_register_snapshot": register_snapshot,
        }
    )
    package = create_pilot_export_package(
        report_text,
        audit_path=audit_path,
        review_record=review_record,
        package_context=package_context,
        register_snapshot=register_snapshot,
        analysis={},
    )

    assert package["filename"].endswith("_pilot_export_package.zip")
    with ZipFile(BytesIO(package["content"])) as archive:
        names = set(archive.namelist())

    assert "governance/package_manifest.json" in names
    assert "governance/reviewer_signoff.json" in names
    assert "governance/data_register.csv" in names
    assert any(name.startswith("reports/cairns_") and name.endswith(".md") for name in names)


def test_report_exporters_create_valid_pdf_and_docx_files():
    pdf_content = create_report_pdf(SAMPLE_REPORT)
    docx_content = create_report_docx(SAMPLE_REPORT)

    assert pdf_content.startswith(b"%PDF")
    assert len(pdf_content) > 1000
    assert docx_content.startswith(b"PK")

    with ZipFile(BytesIO(docx_content)) as document:
        names = set(document.namelist())

    assert "word/document.xml" in names
    assert "docProps/core.xml" in names


def test_governed_export_metadata_uses_report_values_instead_of_structural_title():
    governed_report = """# 1. Title
## Title
Bushfire Preparedness Planning Report for Cairns, Queensland - Draft for Human Review

## Evidence Tables
- U0 User-provided / unverified context: location: Cairns, Queensland; audience: Council resilience officers
"""

    metadata = extract_report_metadata(governed_report)

    assert metadata == {
        "title": "Bushfire Preparedness Planning Report for Cairns, Queensland - Draft for Human Review",
        "location": "Cairns, Queensland",
        "audience": "Council resilience officers",
    }


def test_wide_export_table_becomes_readable_record_layout():
    long_planning_context = " ".join(
        [
            "Confirm official information sources, responsible owners, candidate assembly-point approval,"
            " backup communications, smoke-health support and human review before operational use."
        ]
        * 12
    )
    wide_report = (
        """# Evidence report

- Verify official sources before use.

| Source | Page | Hybrid score | Dense score | BM25 score | Document date | Passage hash | URL |
| --- | --- | --- | --- | --- | --- | --- | --- |
| Official source | web / 1 | 1.0 | 0.55 | 8.30 | 2024-02-26 | abcdef0123456789 | https://example.gov.au/long/source/path |
"""
        + f"""

| Contribution | Current output | Evidence level / note |
| --- | --- | --- |
| Planning priorities | {long_planning_context} | Deterministic planning transformation |
"""
    )

    docx_content = create_report_docx(wide_report)
    pdf_content = create_report_pdf(wide_report)

    document = Document(BytesIO(docx_content))
    table_column_counts = [len(table.columns) for table in document.tables]
    pdf_reader = PdfReader(BytesIO(pdf_content))
    pdf_pages = [page.extract_text() or "" for page in pdf_reader.pages]
    pdf_text = "\n".join(pdf_pages)

    assert max(table_column_counts) <= 4
    assert any(paragraph.text == "Record 1" for paragraph in document.paragraphs)
    assert pdf_text.count("Record 1") == 2
    assert "Official source" in pdf_text
    assert "Planning priorities" in pdf_text
    assert "- Verify official sources before use." in pdf_text
    assert "\x00" not in pdf_text
    assert "\x7f" not in pdf_text
    assert all("BushfireReadyGPT" in page for page in pdf_pages)
    assert all("Planning support only." in page for page in pdf_pages)
    for page in pdf_reader.pages:
        operators = page.get_contents().get_data().split()
        assert operators.count(b"q") == operators.count(b"Q")


def test_pdf_human_signoff_starts_on_a_dedicated_page():
    report = """# Sample report

Preparedness planning content.

## Human Review Sign-off

| Field | Value |
| --- | --- |
| Review status | Draft - human review required |
"""

    pages = [page.extract_text() or "" for page in PdfReader(BytesIO(create_report_pdf(report))).pages]
    signoff_page = next(index for index, text in enumerate(pages) if "Human Review Sign-off" in text)

    assert signoff_page >= 2
    assert "Human Review Sign-off" not in pages[signoff_page - 1]


def test_pilot_export_package_includes_report_formats_and_manifest_boundary(tmp_path, monkeypatch):
    _allow_governed_package_export(monkeypatch)
    review_record = {
        "approval_status": DRAFT_STATUS,
        "reviewer_name": "Test Reviewer",
    }
    report_text = append_human_signoff(SAMPLE_REPORT, review_record)
    package_context = {
        "location": "Cairns, Queensland",
        "report_status": "Draft - human review required",
        "report_id": "report-formats-test",
        "report_version": 1,
    }
    register_snapshot = build_export_register_snapshot()
    monkeypatch.setenv("BUSHFIRE_AUDIT_DIR", str(tmp_path))
    audit_path = save_report_audit(
        {
            "report_id": "report-formats-test",
            "report_version": 1,
            "report_text": report_text,
            "inputs": {"location": "Cairns, Queensland"},
            "human_review": review_record,
            "export_register_snapshot": register_snapshot,
        }
    )
    package = create_pilot_export_package(
        report_text,
        audit_path=audit_path,
        review_record=review_record,
        package_context=package_context,
        register_snapshot=register_snapshot,
        analysis={},
    )

    assert package["manifest"]["package_schema"] == PILOT_EXPORT_SCHEMA
    assert package["manifest"]["governed_quality"] == {
        "version": QUALITY_POLICY_VERSION,
        "fingerprint": QUALITY_POLICY_FINGERPRINT,
        "manifest": package["manifest"]["governed_quality"]["manifest"],
        "approval_gate_passed": True,
        "analysis_sha256": package["manifest"]["governed_quality"]["analysis_sha256"],
        "quality_sha256": package["manifest"]["governed_quality"]["quality_sha256"],
    }
    assert "Not live emergency advice" in package["manifest"]["safety_boundary"]

    with ZipFile(BytesIO(package["content"])) as archive:
        names = set(archive.namelist())

    assert any(name.endswith(".pdf") for name in names)
    assert any(name.endswith(".docx") for name in names)
    assert "governance/package_manifest.json" in names


def test_pilot_export_package_writes_one_complete_manifest_with_audit(tmp_path, monkeypatch):
    _allow_governed_package_export(monkeypatch)
    monkeypatch.setenv("BUSHFIRE_AUDIT_DIR", str(tmp_path))
    review_record = {
        "approval_status": DRAFT_STATUS,
        "reviewer_name": "Test Reviewer",
    }
    report_text = append_human_signoff(SAMPLE_REPORT, review_record)
    register_snapshot = build_export_register_snapshot()
    audit_path = Path(
        save_report_audit(
            {
                "report_id": "export-test",
                "report_version": 1,
                "report_text": report_text,
                "inputs": {"location": "Cairns, Queensland"},
                "human_review": review_record,
                "export_register_snapshot": register_snapshot,
            }
        )
    )

    package = create_pilot_export_package(
        report_text,
        audit_path=audit_path,
        review_record=review_record,
        package_context={
            "location": "Cairns, Queensland",
            "report_status": "Draft - human review required",
            "report_id": "export-test",
            "report_version": 1,
        },
        register_snapshot=register_snapshot,
        analysis={},
    )

    with ZipFile(BytesIO(package["content"])) as archive:
        names = archive.namelist()
        manifest = json.loads(archive.read("governance/package_manifest.json"))

    assert names.count("governance/package_manifest.json") == 1
    assert names.count("governance/audit_record.json") == 1
    assert manifest["included_files"].count("governance/audit_record.json") == 1
    assert manifest == package["manifest"]


def test_ollama_connection_error_has_safe_recovery_guidance():
    error = APIConnectionError(request=httpx.Request("POST", "http://localhost:11434/v1/chat/completions"))

    message = model_service_error_message(error, provider="ollama", model_name="qwen2.5:7b")

    assert "Cannot reach the local Ollama service" in message
    assert "ollama serve" in message
    assert "Traceback" not in message


def test_ollama_missing_model_error_has_pull_command():
    request = httpx.Request("POST", "http://localhost:11434/v1/chat/completions")
    response = httpx.Response(404, request=request)
    error = APIStatusError("model not found", response=response, body={"error": "model not found"})

    message = model_service_error_message(error, provider="ollama", model_name="qwen2.5:7b")

    assert "configured model `qwen2.5:7b` is unavailable" in message
    assert "ollama pull qwen2.5:7b" in message
    assert "model not found" not in message


def test_report_validation_requires_human_review_details_for_approved_outputs():
    base_inputs = {
        "location": "Cairns, Queensland",
        "audience": "students and teachers",
        "concerns": ["Evacuation"],
        "report_status": "Draft - human review required",
        "organisation_name": "",
        "reviewer_name": "",
        "reviewer_role": "",
    }

    assert validate_report_inputs(base_inputs) is None

    approved_inputs = dict(base_inputs)
    approved_inputs["report_status"] = "Approved by organisation"
    assert "reviewer name" in validate_report_inputs(approved_inputs)

    approved_inputs.update(
        {
            "organisation_name": "Cairns Council",
            "reviewer_name": "Test Reviewer",
            "reviewer_role": "Preparedness officer",
        }
    )
    assert validate_report_inputs(approved_inputs) is None


def test_report_validation_rejects_missing_required_form_fields():
    base_inputs = {
        "location": "",
        "audience": "students and teachers",
        "concerns": ["Evacuation"],
        "report_status": "Draft - human review required",
    }

    assert "location and audience" in validate_report_inputs(base_inputs)

    base_inputs["location"] = "Cairns, Queensland"
    base_inputs["audience"] = ""
    assert "location and audience" in validate_report_inputs(base_inputs)

    base_inputs["audience"] = "students and teachers"
    base_inputs["concerns"] = []
    assert "at least one focus area" in validate_report_inputs(base_inputs)


def test_quality_agent_accepts_markdown_checkboxes_and_human_review_boundary():
    report = """
    # Cairns Preparedness Report

    ## Executive Summary
    ## Purpose and Scope
    ## Selected Geography
    ## Data Sources
    ## Local Risk Context
    ## Evacuation
    ## Candidate Assembly Point Criteria
    ## Roles and Responsibilities
    ## Communication
    ## First Aid
    ## Action Plan
    Day 1 actions are listed below.
    ## Human Review and Approval Checklist
    - [ ] Confirm official source checks.
    ## Evidence Tables
    Selected geography, community indicators, ASGS and official source register are included.
    ### Evidence Confidence and Provenance
    O1 official reference, P2 processed data, R3 rule inference, A4 AI draft and U0 unverified input.
    ## Safety Disclaimer
    This is not live emergency advice. Official evacuation order information must come from
    Queensland Fire Department, QFES, Queensland Disaster, Cairns Regional Council,
    Bureau of Meteorology / BoM and 000.
    ## Human Review Sign-off
    This report remains a draft until approved by a responsible organisation.
    """

    quality = ReportQualityAgent().run(report)
    checks = {item["name"]: item["status"] for item in quality["checks"]}

    assert checks["Checklist"] == "pass"
    assert checks["Evidence confidence"] == "pass"
    assert checks["Human review status"] == "pass"


@pytest.mark.parametrize(
    "first_action",
    [
        "Day One: confirm responsible owners.",
        "The first day assigns the preparedness coordinator.",
        "Immediate action: verify official contacts.",
        "Complete within the first 24 hrs: confirm the contact tree.",
        "0–24 hours | Responsible lead | Verify warning channels",
        "| 1 | Responsible lead | Verify warning channels |",
    ],
)
def test_quality_agent_accepts_equivalent_immediate_action_labels(first_action):
    result = ReportQualityAgent()._check_action_plan(f"## Action Plan\n{first_action}")

    assert result["status"] == "pass"


def test_quality_agent_does_not_treat_week_one_as_an_immediate_action():
    result = ReportQualityAgent()._check_action_plan("## Action Plan\nWeek 1: consider assigning an owner.")

    assert result["status"] == "fail"


def test_quality_agent_counts_nested_subsection_content_under_required_heading():
    sections = ReportQualityAgent()._extract_sections(
        """## 8. Evacuation Planning
### Warning Monitoring
Monitor official warnings and nominate a responsible officer before each review cycle.
### Accountability
Maintain a locally approved roll-call process and record unresolved gaps.
## 9. Candidate Assembly Point Criteria
Use criteria that require local approval before a candidate location is treated as suitable.
"""
    )

    assert "official warnings" in sections["evacuation planning"]
    assert "roll-call process" in sections["evacuation planning"]
    assert "local approval" in sections["candidate assembly point criteria"]


def test_repair_prompt_omits_full_previous_response_to_protect_local_context_window():
    previous = "PRIVATE-INCOMPLETE-DRAFT " * 1000
    prompt = build_report_repair_prompt(
        "Original governed request",
        previous,
        {"approval_gate": {"blocking_failures": [{"name": "Required sections", "detail": "missing"}]}},
    )

    assert "Original governed request" in prompt
    assert "PRIVATE-INCOMPLETE-DRAFT" not in prompt
    assert f"previous {len(previous)}-character response" in prompt


def test_repair_prompt_gives_exhaustive_premises_status_rewrite_guidance():
    prompt = build_report_repair_prompt(
        "Original governed request",
        "The school is available.",
        {
            "approval_gate": {
                "blocking_failures": [
                    {
                        "name": "Safety boundary assertions",
                        "detail": "Remove prohibited operational assertions (premises_status_assertion).",
                    }
                ]
            }
        },
    )

    assert "PLACE/PREMISES REWRITE" in prompt
    assert "safe, open, approved, authorised, available, operational, suitable or cleared" in prompt
    assert "unverified candidate pending current verification" in prompt
    assert "prose, tables, checklists and examples" in prompt
    assert "The school is available." not in prompt


def test_generated_checklist_bullets_are_normalized_without_changing_other_sections():
    narrative = """## 14. Human Review and Approval Checklist
- Confirm geography.
* Confirm official sources.
1. Record approval status.
## 15. Safety Disclaimer
- This ordinary bullet must stay ordinary.
"""

    normalized = normalize_generated_narrative(narrative)

    assert "- [ ] Confirm geography." in normalized
    assert "- [ ] Confirm official sources." in normalized
    assert "- [ ] Record approval status." in normalized
    assert "- This ordinary bullet must stay ordinary." in normalized


def test_generated_quality_blocks_missing_rag_source_attribution():
    quality = assess_generated_narrative(
        "## Data Sources and Limitations\nOfficial sources must be verified.",
        {"knowledge": {"retrieved_chunks": [{"title": "Your Bushfire Plan", "agency": "Country Fire Authority"}]}},
    )

    failures = quality["approval_gate"]["blocking_failures"]
    assert any(item["name"] == "RAG source attribution" for item in failures)


def test_canonical_gate_does_not_count_deterministic_appendix_as_narrative_attribution():
    analysis = {
        "knowledge": {
            "retrieved_chunks": [
                {
                    "source_id": "qld_prepare_home",
                    "title": "Prepare your home for bushfire season",
                    "agency": "Queensland Government",
                }
            ]
        }
    }
    report = append_human_signoff(
        append_evidence_tables(
            apply_governance_notice("## Data Sources and Limitations\nOfficial sources require human review."),
            analysis,
        ),
        {},
    )

    quality = evaluate_governed_report(report, analysis)

    rag_check = next(item for item in quality["checks"] if item["name"] == "RAG source attribution")
    assert rag_check["status"] == "fail"


def test_rag_source_attribution_rejects_agency_acronym_without_canonical_label():
    attributed = attributed_rag_source_ids(
        "## Data Sources and Limitations\nThe draft uses DFES guidance and requires local verification.",
        [
            {
                "source_id": "wa_prepare_bushfire",
                "title": "Prepare for a bushfire",
                "agency": "Department of Fire and Emergency Services Western Australia",
            }
        ],
    )

    assert attributed == set()


@pytest.mark.parametrize(
    ("claim", "expected"),
    [
        ("Preparedness plans should document responsibilities before fire season.", {"wa_prepare_bushfire"}),
        ("Preparedness plans need review", set()),
        ("Brief plan.", set()),
        ("", set()),
    ],
)
def test_rag_source_attribution_requires_a_substantive_complete_sentence_ending_in_label(claim, expected):
    chunks = [
        {
            "source_id": "wa_prepare_bushfire",
            "title": "Prepare for a bushfire",
            "agency": "Department of Fire and Emergency Services Western Australia",
        }
    ]
    label = format_rag_attribution(chunks[0])
    narrative = f"## Data Sources and Limitations\n{claim} {label}".rstrip()

    assert attributed_rag_source_ids(narrative, chunks) == expected


def test_rag_source_attribution_rejects_citation_before_more_sentence_text_or_outside_source_section():
    chunks = [
        {
            "source_id": "wa_prepare_bushfire",
            "title": "Prepare for a bushfire",
            "agency": "Department of Fire and Emergency Services Western Australia",
        }
    ]
    label = format_rag_attribution(chunks[0])
    claim = "Preparedness plans should document responsibilities before fire season."

    assert (
        attributed_rag_source_ids(f"## Data Sources and Limitations\n{claim} {label} More commentary.", chunks) == set()
    )
    assert attributed_rag_source_ids(f"## Preparedness Priorities\n{claim} {label}", chunks) == set()


def test_rag_source_attribution_handles_missing_source_id_without_crashing():
    quality = assess_generated_narrative(
        "## Data Sources and Limitations\n[O1-RAG][source_id=unknown-source] Untitled official source",
        {"knowledge": {"retrieved_chunks": [{"title": "Incomplete source", "agency": "Test agency"}]}},
    )

    check = next(item for item in quality["checks"] if item["name"] == "RAG source attribution")
    assert check["status"] == "fail"


def test_rag_source_attribution_ignores_fake_section_inside_fenced_code():
    chunks = [{"source_id": "qld-guide", "title": "Queensland Guide", "agency": "Test agency"}]
    narrative = """## Preparedness Priorities
```markdown
## Data Sources and Limitations
[O1-RAG][source_id=qld-guide] Queensland Guide
```
## Action Plan
Confirm real sources separately.
"""

    assert attributed_rag_source_ids(narrative, chunks) == set()


@pytest.fixture
def registered_official_sources():
    return [
        {"id": "qfd_qfes", "name": "Queensland Fire Department / QFES"},
        {"id": "cairns_disaster", "name": "Cairns Regional Council Disaster Information"},
        {"id": "bom_qld_warnings", "name": "Bureau of Meteorology Queensland Warnings"},
    ]


@pytest.mark.parametrize("list_marker", ["", "- ", "* ", "1. "])
def test_official_source_gate_accepts_each_registered_label_only_on_its_own_plain_or_list_line(
    registered_official_sources, list_marker
):
    first, second = registered_official_sources[:2]
    narrative = (
        "## Data Sources and Limitations\n"
        f"{list_marker}{format_official_attribution(first)}\n"
        f"{list_marker}{format_official_attribution(second)}"
    )

    assert ReportQualityAgent()._check_official_sources(narrative, registered_official_sources)["status"] == "pass"


def test_official_source_gate_requires_two_exact_registered_labels_in_source_section(registered_official_sources):
    first = registered_official_sources[0]
    narrative = f"## Data Sources and Limitations\n{format_official_attribution(first)}"

    assert ReportQualityAgent()._check_official_sources(narrative, registered_official_sources)["status"] == "fail"


@pytest.mark.parametrize(
    "hidden_lines",
    [
        (
            "This sentence cites [O1][source_id=qfd_qfes] Queensland Fire Department / QFES and "
            "[O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information."
        ),
        (
            "[qfd]: [O1][source_id=qfd_qfes] Queensland Fire Department / QFES\n"
            "[council]: [O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information"
        ),
        (
            "<!-- [O1][source_id=qfd_qfes] Queensland Fire Department / QFES -->\n"
            "<!-- [O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information -->"
        ),
        (
            "<!-- [O1][source_id=qfd_qfes] Queensland Fire Department / QFES\n"
            "[O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information"
        ),
        (
            "<div>[O1][source_id=qfd_qfes] Queensland Fire Department / QFES</div>\n"
            "<div>[O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information</div>"
        ),
        (
            "<div>[O1][source_id=qfd_qfes] Queensland Fire Department / QFES\n"
            "[O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information"
        ),
        (
            "<span hidden>[O1][source_id=qfd_qfes] Queensland Fire Department / QFES</span>\n"
            '<span style="display:none">[O1][source_id=cairns_disaster] '
            "Cairns Regional Council Disaster Information</span>"
        ),
        (
            "<template>[O1][source_id=qfd_qfes] Queensland Fire Department / QFES\n"
            "[O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information"
        ),
        (
            "```text\n[O1][source_id=qfd_qfes] Queensland Fire Department / QFES\n"
            "[O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information\n```"
        ),
        (
            "`[O1][source_id=qfd_qfes] Queensland Fire Department / QFES`\n"
            "    [O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information"
        ),
    ],
    ids=[
        "sentence",
        "reference-definition",
        "closed-comment",
        "unclosed-comment",
        "closed-html",
        "unclosed-html",
        "hidden-html",
        "unclosed-non-visible-html",
        "fenced-code",
        "inline-and-indented-code",
    ],
)
def test_official_source_gate_rejects_labels_that_are_not_visible_standalone_lines(
    registered_official_sources, hidden_lines
):
    narrative = f"## Data Sources and Limitations\n{hidden_lines}"

    assert ReportQualityAgent()._check_official_sources(narrative, registered_official_sources)["status"] == "fail"


@pytest.mark.parametrize(
    "opening_tag",
    ["<div hidden>", "<div style=display&#58;none>", "<div\n style=display&#58;none>"],
    ids=["hidden-attribute", "entity-css", "multiline-tag"],
)
def test_source_gates_reject_an_entire_source_section_inside_a_hidden_container(
    registered_official_sources, opening_tag
):
    first, second = registered_official_sources[:2]
    rag_source = {"source_id": "qld-guide", "title": "Queensland Guide"}
    narrative = (
        f"{opening_tag}\n"
        "## Data Sources and Limitations\n"
        f"- {format_official_attribution(first)}\n"
        f"- {format_official_attribution(second)}\n"
        "Preparedness plans should be reviewed before each season. "
        f"{format_rag_attribution(rag_source)}\n"
        "</div>"
    )

    assert ReportQualityAgent()._check_official_sources(narrative, registered_official_sources)["status"] == "fail"
    assert attributed_rag_source_ids(narrative, [rag_source]) == set()


def test_official_source_gate_rejects_unknown_or_fenced_labels():
    sources = [
        {"id": "qfd_qfes", "name": "Queensland Fire Department / QFES"},
        {"id": "cairns_disaster", "name": "Cairns Regional Council Disaster Information"},
    ]
    narrative = """## Preparedness Priorities
```markdown
## Data Sources and Limitations
[O1][source_id=qfd_qfes] Queensland Fire Department / QFES
[O1][source_id=cairns_disaster] Cairns Regional Council Disaster Information
```
## Data Sources and Limitations
[O1][source_id=unknown] Queensland Fire Department / QFES
[O1][source_id=cairns_disaster] Cairns Regional Council
"""

    assert ReportQualityAgent()._check_official_sources(narrative, sources)["status"] == "fail"


def test_operational_looking_registered_title_is_removed_from_safety_lint_but_ordinary_prose_is_not():
    source = {"id": "road_status_register", "name": "M1 road is open"}
    label_only = f"## Data Sources and Limitations\n{format_official_attribution(source)}"
    label_result = ReportQualityAgent().run(label_only, official_sources=[source])
    prose_result = ReportQualityAgent().run(
        f"{label_only}\n## Local Risk Context\nThe M1 road is open.",
        official_sources=[source],
    )

    label_safety = next(item for item in label_result["checks"] if item["name"] == "Safety boundary assertions")
    prose_safety = next(item for item in prose_result["checks"] if item["name"] == "Safety boundary assertions")
    assert label_safety["status"] == "pass"
    assert prose_safety["status"] == "fail"


def test_safety_qualifier_on_previous_line_does_not_suppress_later_operational_assertion():
    evaluation = evaluate_safety_boundaries("Do not claim or infer operational road status.\nThe M1 road is open.")

    assert evaluation["passed"] is False
    assert {item["code"] for item in evaluation["violations"]} == {"road_status_assertion"}


@pytest.mark.parametrize(
    "obfuscated_claim",
    [
        "Smith R&#111;ad is open.",
        "Smith Ro\u200bad is open.",
        "Smith Ro\ufe0fad is open.",
        "Smith Ro\u0301ad is open.",
        "Ｓｍｉｔｈ Ｒｏａｄ ｉｓ ｏｐｅｎ．",
    ],
    ids=["html-entity", "zero-width", "variation-selector", "combining-mark", "nfkc-fullwidth"],
)
def test_report_quality_safety_lint_normalizes_render_equivalent_operational_claims(obfuscated_claim):
    result = ReportQualityAgent().run(SAMPLE_REPORT + "\n\n" + obfuscated_claim)
    check = next(item for item in result["checks"] if item["name"] == "Safety boundary assertions")

    assert check["status"] == "fail"


@pytest.mark.parametrize(
    "model_markup",
    [
        "See https://attacker.example/path for details.",
        "Use [this link](//attacker.example/path) for details.",
        "Use [this link](//127.0.0.1:8080/admin) for details.",
        "Use [this link](//[::1]:8080/admin) for details.",
        "Use [this link](//localhost:8080/admin) for details.",
        "Use [this link](https&#58;//attacker.example/path) for details.",
        "Use [this link](ftp://attacker.example/file) for details.",
        "Use [this link](mailto:attacker@example.test) for details.",
        "Use [this link](data:text/html,unsafe) for details.",
        "Use [this link](javascript:alert(1)) for details.",
        "[external]: relative/or/custom-scheme-target",
        "<div>Unsafe hidden narrative content.</div>",
        "<div\n hidden>Unsafe hidden narrative content.</div>",
        "<!-- hidden narrative -->",
    ],
    ids=[
        "https",
        "scheme-relative-domain",
        "scheme-relative-ipv4",
        "scheme-relative-ipv6",
        "scheme-relative-localhost",
        "entity-url",
        "ftp",
        "mailto",
        "data",
        "javascript",
        "reference-link",
        "html",
        "multiline-html",
        "html-comment",
    ],
)
def test_report_quality_blocks_model_authored_links_and_raw_html(model_markup):
    result = ReportQualityAgent().run(SAMPLE_REPORT + "\n\n" + model_markup)
    failures = {item["name"] for item in result["approval_gate"]["blocking_failures"]}

    assert failures & {"Model-authored URLs", "Model-authored raw HTML"}


def test_structural_gate_is_independent_from_safety_and_rag_failures():
    quality = {
        "checks": [
            {"name": "Required sections", "status": "pass"},
            {"name": "Safety boundary assertions", "status": "fail"},
            {"name": "RAG source attribution", "status": "fail"},
        ]
    }

    assert structural_gate_passed(quality) is True
    quality["checks"][0]["status"] = "fail"
    assert structural_gate_passed(quality) is False


def test_incomplete_report_body_fails_quality_and_cannot_be_approved():
    report = apply_governance_notice("")
    report = append_evidence_tables(report, {})
    report = append_human_signoff(report, {})
    quality = ReportQualityAgent().run(report)
    checks = {item["name"]: item["status"] for item in quality["checks"]}
    review_record = {
        "approval_status": "Approved by organisation",
        "reviewer_name": "Test Reviewer",
        "reviewer_role": "Safety reviewer",
        "organisation_name": "Test Organisation",
        "review_date": "2025-01-01",
        "review_checklist": build_review_checklist_snapshot(lambda _item_id: True),
        "review_checklist_complete": True,
    }

    assert checks["Required sections"] == "fail"
    assert quality["summary"]["failed"] >= 1
    assert (
        "failed governed"
        in validate_review_record(
            review_record,
            quality,
            {
                "text": report,
                "analysis": {"data_integrity": {"core_ready": True, "custom_data": False}},
            },
        ).lower()
    )


def test_keyword_stuffing_cannot_satisfy_the_narrative_quality_gate():
    stuffed = " ".join(
        [
            "Executive Summary Purpose and Scope Selected Geography and Key Assumptions",
            "Data Sources and Limitations Local Risk Context Preparedness Priorities",
            "Evacuation Planning Candidate Assembly Point Criteria Roles and Responsibilities",
            "Communication and Inclusion Needs First Aid Training and Exercises Action Plan",
            "Human Review and Approval Checklist Safety Disclaimer official ASGS Day 1 000",
        ]
    )
    report = append_human_signoff(append_evidence_tables(apply_governance_notice(stuffed), {}), {})

    quality = ReportQualityAgent().run(report)
    checks = {item["name"]: item["status"] for item in quality["checks"]}

    assert checks["Substantive narrative"] == "fail"
    assert checks["Required sections"] == "fail"
    assert quality["approval_gate"]["passed"] is False


def test_candidate_assembly_quality_check_distinguishes_negation_from_confirmation():
    agent = ReportQualityAgent()

    safe = agent._check_candidate_assembly_language(
        "The gymnasium is a candidate only and is not confirmed safe until local approval."
    )
    unsafe = agent._check_candidate_assembly_language("The gymnasium is a confirmed safe assembly point.")

    assert safe["status"] == "pass"
    assert unsafe["status"] == "fail"


def test_report_quality_gate_blocks_live_or_unverified_safety_assertions():
    unsafe = ReportQualityAgent().run(
        SAMPLE_REPORT + "\n\nThe highway remains open. Identify safe evacuation routes for residents."
    )
    check = next(item for item in unsafe["checks"] if item["name"] == "Safety boundary assertions")

    assert check["status"] == "fail"
    assert unsafe["approval_gate"]["passed"] is False


def test_report_quality_gate_allows_candidate_routes_with_explicit_verification():
    result = ReportQualityAgent()._check_safety_boundaries(
        "Map multiple candidate evacuation routes and verify them with local authorities. "
        "No evacuation route is confirmed safe."
    )

    assert result["status"] == "pass"


def test_repeated_filler_under_every_required_heading_is_blocked():
    narrative = "\n".join(
        f"## {heading}\n" + "planning " * 30 for heading in ReportQualityAgent.REQUIRED_SECTION_HEADINGS
    )
    report = append_human_signoff(
        append_evidence_tables(apply_governance_notice(narrative), {}),
        {},
    )
    quality = ReportQualityAgent().run(report)
    checks = {item["name"]: item["status"] for item in quality["checks"]}

    assert checks["Substantive narrative"] == "fail"
    assert checks["Required sections"] == "fail"
    assert quality["approval_gate"]["passed"] is False


def test_approval_gate_accepts_a_completed_review_when_exact_report_passes():
    review_record = {
        "approval_status": "Approved by organisation",
        "reviewer_name": "Test Reviewer",
        "reviewer_role": "Safety reviewer",
        "organisation_name": "Test Organisation",
        "review_date": "2025-01-01",
        "review_checklist": build_review_checklist_snapshot(lambda _item_id: True),
        "review_checklist_complete": True,
    }
    narrative = dedent("""
    # Cairns Preparedness Report

    ## Executive Summary
    This draft supports local preparedness planning and requires responsible human review.
    It helps the organisation prepare people, information and facilities before a bushfire or
    smoke event. It does not predict fire behaviour. Local leaders should verify every proposed
    action against current conditions, official advice and the needs of the people who may rely
    on the plan. Findings should be discussed with staff, community representatives and relevant
    emergency-management partners before they are adopted.
    ## Purpose and Scope
    It covers planning before an incident and is not live emergency advice.
    ## Selected Geography and Key Assumptions
    Cairns, Queensland is the selected planning area; local conditions must be verified.
    ## Data Sources and Limitations
    Official sources, processed community data, and user context have different confidence levels.
    [O1][source_id=qfd_qfes] Queensland Fire Department / QFES
    [O1][source_id=bom_qld_warnings] Bureau of Meteorology Queensland Warnings
    The Bureau of Meteorology, the state fire service, local council and other emergency services
    remain authoritative for current warnings. Processed indicators are planning context rather
    than facts about an individual person. User-provided details are unverified until a responsible
    reviewer checks organisational records. Call 000 when life or property is in immediate danger.
    ## Local Risk Context
    Seasonal bushfire exposure and smoke impacts require local validation.
    ## Preparedness Priorities
    Verify contacts, routes, accessibility needs, and current official warnings.
    Maintain alternatives for power, telecommunications and transport disruption. Record who owns
    each action, when it is due, and what evidence will demonstrate completion. Recheck the plan
    after exercises, staffing changes and material changes to local risk.
    ## Evacuation Planning
    Follow emergency-service directions; do not infer that any route is safe.
    Identify more than one candidate route and document the authority responsible for deciding
    whether movement is appropriate. Plans must account for mobility, language, supervision,
    transport and reunification needs without promising that a route will remain available.
    ## Candidate Assembly Point Criteria
    Responsible authorities must assess hazards, access, capacity, and alternatives.
    A gymnasium, library, sports field or carpark may only be recorded as a candidate pending a
    site-specific inspection and official advice. Document shade, smoke exposure, accessibility,
    communications, first aid access, capacity and a backup location.
    ## Roles and Responsibilities
    Coordinators, communications leads, first aid staff, and wardens require named backups.
    ## Communication and Inclusion Needs
    Use accessible, redundant channels and support people with additional needs.
    ## First Aid, Training and Exercises
    Check supplies and run a documented exercise with corrective actions.
    ## Action Plan
    Day 1: verify official contacts. Day 2: inspect routes. Day 3: exercise the plan.
    ## Human Review and Approval Checklist
    - [ ] Confirm official sources and current local arrangements.
    ## Safety Disclaimer
    This report is not live emergency advice or an evacuation order. Follow current official
    emergency-service instructions and call 000 if life or property is in immediate danger.
    """)
    report = append_human_signoff(
        append_evidence_tables(apply_governance_notice(narrative), {}),
        review_record,
    )
    quality = ReportQualityAgent().run(report)

    assert quality["approval_gate"]["passed"] is True, quality["approval_gate"]
    assert (
        validate_review_record(
            review_record,
            quality,
            {
                "text": report,
                "analysis": {
                    "data": {
                        "sources": [
                            {"id": "qfd_qfes", "name": "Queensland Fire Department / QFES"},
                            {
                                "id": "bom_qld_warnings",
                                "name": "Bureau of Meteorology Queensland Warnings",
                            },
                        ]
                    },
                    "data_integrity": {"core_ready": True, "custom_data": False},
                },
            },
        )
        is None
    )


def test_approval_fails_closed_for_missing_or_legacy_quality_results():
    review_record = {
        "approval_status": "Approved by organisation",
        "reviewer_name": "Test Reviewer",
        "reviewer_role": "Safety reviewer",
        "organisation_name": "Test Organisation",
        "review_date": "2025-01-01",
        "review_checklist": build_review_checklist_snapshot(lambda _item_id: True),
        "review_checklist_complete": True,
    }
    report_record = {
        "text": "Governed report",
        "analysis": {"data_integrity": {"core_ready": True, "custom_data": False}},
    }

    assert "failed Governed" in validate_review_record(review_record, None, report_record)
    legacy_quality = {"summary": {"passed": 11, "warnings": 0, "failed": 0, "total": 11}}
    assert "failed Governed" in validate_review_record(review_record, legacy_quality, report_record)


def test_approval_is_blocked_for_unverified_custom_data():
    review_record = {
        "approval_status": "Approved by organisation",
        "reviewer_name": "Test Reviewer",
        "reviewer_role": "Safety reviewer",
        "organisation_name": "Test Organisation",
        "review_date": "2025-01-01",
        "review_checklist": build_review_checklist_snapshot(lambda _item_id: True),
    }
    report_record = {
        "text": "Governed report",
        "analysis": {
            "data_integrity": {
                "core_ready": True,
                "custom_data": True,
                "integrity_status": "Unverified custom data",
            }
        },
    }

    assert (
        "unverified custom data"
        in validate_review_record(
            review_record,
            {},
            report_record,
        ).lower()
    )


def test_approval_is_blocked_for_unverified_selected_map_bundle():
    review_record = {
        "approval_status": "Approved by organisation",
        "reviewer_name": "Test Reviewer",
        "reviewer_role": "Safety reviewer",
        "organisation_name": "Test Organisation",
        "review_date": "2025-01-01",
        "review_checklist": build_review_checklist_snapshot(lambda _item_id: True),
    }
    report_record = {
        "text": "Governed report",
        "area_selection": {
            "state": "Queensland",
            "level": "SA2",
            "area_name": "Cairns City",
        },
        "analysis": {
            "data_integrity": {
                "core_ready": True,
                "custom_data": False,
                "optional_map_state": "present_unverified",
            }
        },
    }

    error = validate_review_record(review_record, {}, report_record)

    assert "sidecar-verified" in error


def test_human_signoff_reflects_completed_review_checklist():
    partial_items = build_review_checklist_snapshot(lambda item_id: item_id == "official_sources")
    incomplete = build_human_signoff({"review_checklist": partial_items, "review_checklist_complete": False})
    complete = build_human_signoff(
        {
            "review_checklist": build_review_checklist_snapshot(lambda _item_id: True),
            "review_checklist_complete": True,
        }
    )

    assert incomplete.count("- [ ]") == 4
    assert incomplete.count("- [x]") == 1
    assert complete.count("- [x]") == 5


def test_governance_notice_satisfies_safety_quality_check():
    report = apply_governance_notice(
        """
        ## Executive Summary
        ## Purpose and Scope
        ## Selected Geography
        ## Data Sources
        ## Local Risk Context
        ## Evacuation
        ## Candidate Assembly Point Criteria
        ## Roles and Responsibilities
        Teacher, student, first aid and communication roles are included.
        ## Communication
        ## First Aid
        ## Action Plan
        Day 1 actions are included.
        ## Human Review and Approval Checklist
        - [ ] Confirm official sources.
        ## Evidence Tables
        Selected geography, community indicators, ASGS and official source register are included.
        ## Human Review Sign-off
        This remains a draft.
        State fire service, local council, Bureau of Meteorology BoM, official emergency services and 000.
        """
    )

    quality = ReportQualityAgent().run(report)
    checks = {item["name"]: item["status"] for item in quality["checks"]}

    assert checks["Safety disclaimer"] == "pass"


def test_official_sources_are_state_specific_for_non_queensland_locations():
    cases = [
        ("Sydney, NSW", "New South Wales", "NSW Rural Fire Service"),
        ("Melbourne, Victoria", "Victoria", "VicEmergency"),
        ("Perth Hills, WA", "Western Australia", "Emergency WA"),
    ]

    for location, expected_state, expected_source_name in cases:
        analysis = run_analysis_pipeline(
            location=location,
            audience="community residents",
            scenario="Community workshop material",
            concerns=["Official information sources", "Evacuation"],
            timeframe="7-day action plan",
            extra_context="Cross-state official source selection test.",
        )

        source_names = [source.get("name", "") for source in analysis["data"].get("sources", [])]
        source_text = " | ".join(source_names)

        assert analysis["profile"]["state"] == expected_state
        assert expected_source_name in source_text
        assert "Triple Zero" in source_text
        assert "Queensland Fire" not in source_text
        assert "Cairns Regional Council" not in source_text


def test_cairns_still_receives_queensland_and_local_sources():
    analysis = run_analysis_pipeline(
        location="Cairns, Queensland",
        audience="council officers",
        scenario="Council community preparedness",
        concerns=["Official information sources", "Evacuation"],
        timeframe="7-day action plan",
        extra_context="Queensland source selection test.",
    )

    source_names = [source.get("name", "") for source in analysis["data"].get("sources", [])]
    source_text = " | ".join(source_names)

    assert analysis["profile"]["state"] == "Queensland"
    assert "Queensland Fire" in source_text
    assert "Cairns Regional Council" in source_text
    assert "Triple Zero" in source_text


def test_state_abbreviations_require_word_boundaries():
    profile = ProfileAgent().run(
        "Wagga Wagga",
        "community residents",
        "Community preparedness",
        ["Evacuation"],
        "7-day action plan",
        "",
    )

    assert profile["state"] == "New South Wales"


def test_risk_context_uses_resolved_state_without_cross_state_abbreviation_matches():
    cases = [
        (
            "Wagga Wagga, New South Wales",
            "new_south_wales_general",
            {"western_australia_general"},
        ),
        (
            "Mount Isa, Queensland",
            "queensland_general",
            {"south_australia_general", "northern_territory_general"},
        ),
        (
            "Cairns",
            "queensland_general",
            set(),
        ),
        (
            "Townsville",
            "queensland_general",
            set(),
        ),
        (
            "Albany, Queensland",
            "queensland_general",
            {"western_australia_general"},
        ),
    ]

    for location, expected_rule, forbidden_rules in cases:
        analysis = run_analysis_pipeline(
            location=location,
            audience="community residents",
            scenario="Household bushfire preparedness",
            concerns=["Evacuation"],
            timeframe="7-day action plan",
            extra_context="State rule matching regression test.",
        )
        matched_rules = set(analysis["risk_context"]["matched_rule_ids"])

        assert expected_rule in matched_rules
        assert matched_rules.isdisjoint(forbidden_rules)


def test_remote_scenario_does_not_borrow_an_unrelated_same_state_profile():
    analysis = run_analysis_pipeline(
        location="Longreach, Queensland",
        audience="remote community residents",
        scenario="Remote community preparedness",
        concerns=["Evacuation", "Road disruption"],
        timeframe="7-day action plan",
        extra_context="Remote profile fallback regression test.",
    )

    assert analysis["community"]["matched_location"] is None
    assert analysis["community"]["indicators"] == {}
    assert "No local community profile row matched" in analysis["community"]["vulnerability_notes"][0]


def test_configured_remote_demo_profile_still_matches_exact_location():
    analysis = run_analysis_pipeline(
        location="Remote Queensland Community, Queensland",
        audience="remote community residents",
        scenario="Remote community preparedness",
        concerns=["Evacuation", "Road disruption"],
        timeframe="7-day action plan",
        extra_context="Configured remote profile regression test.",
    )

    assert analysis["community"]["matched_location"] == "Remote Queensland Community, Queensland"


def test_explicit_area_selection_overrides_location_profile(monkeypatch):
    agent = CommunityVulnerabilityAgent()
    monkeypatch.setattr(
        agent,
        "_load_all_sa2_profiles",
        lambda: [
            {
                "sa2_name": "Longreach",
                "sa3_name": "Outback - North",
                "sa4_name": "Queensland - Outback",
                "state_name": "Queensland",
                "population": "2500",
                "older_people_count": "500",
                "language_other_than_english_count": "125",
            }
        ],
    )
    profile = ProfileAgent().run(
        "Cairns, Queensland",
        "community residents",
        "Community preparedness",
        ["Evacuation"],
        "7-day action plan",
        "",
    )

    result = agent.run(
        profile,
        area_selection={"level": "SA2", "area_name": "Longreach", "state": "Queensland"},
    )

    assert result["matched_location"] == "Longreach, Queensland"
    assert result["indicators"]["population"] == "2500"


def test_invalid_explicit_area_selection_fails_closed_without_location_fallback(monkeypatch):
    agent = CommunityVulnerabilityAgent()
    monkeypatch.setattr(agent, "_load_all_sa2_profiles", lambda: [])
    profile = ProfileAgent().run(
        "Cairns, Queensland",
        "community residents",
        "Community preparedness",
        ["Evacuation"],
        "7-day action plan",
        "",
    )

    result = agent.run(
        profile,
        area_selection={"level": "SA2", "area_name": "Missing Area", "state": "Queensland"},
    )

    assert result["matched_location"] is None
    assert result["indicators"] == {}
    assert "No location-based fallback was used" in result["vulnerability_notes"][0]


def test_exact_cairns_location_keeps_local_official_source():
    analysis = run_analysis_pipeline(
        location="Cairns",
        audience="school staff",
        scenario="School bushfire preparedness",
        concerns=["Official information sources"],
        timeframe="7-day action plan",
        extra_context="",
    )

    source_names = [source.get("name", "") for source in analysis["data"].get("sources", [])]
    assert any("Cairns Regional Council" in name for name in source_names)


def test_geography_validation_rejects_cross_state_map_selection():
    inputs = {
        "location": "Hobart, Tasmania",
        "audience": "community residents",
        "scenario": "Community preparedness",
        "concerns": ["Evacuation"],
        "timeframe": "7-day action plan",
        "extra_context": "",
    }

    error = validate_geography_consistency(
        inputs,
        {"state": "Queensland", "level": "SA4", "area_name": "Cairns"},
    )

    assert "Tasmania" in error
    assert "Queensland" in error
    assert (
        validate_geography_consistency(
            inputs,
            {"state": "Tasmania", "level": "SA4", "area_name": "Hobart"},
        )
        is None
    )


def test_approved_review_requires_identity_fields_and_completed_checklist():
    record = {
        "approval_status": "Approved by organisation",
        "organisation_name": "Cairns Council",
        "reviewer_name": "Test Reviewer",
        "reviewer_role": "Preparedness officer",
        "review_date": "2025-01-01",
        "review_checklist": build_review_checklist_snapshot(),
        "review_checklist_complete": False,
    }
    report_record = {
        "text": "Governed report",
        "analysis": {"data_integrity": {"core_ready": True, "custom_data": False}},
    }

    assert "Complete every Human Review Checklist" in validate_review_record(
        record,
        report_record=report_record,
    )
    record["review_checklist"] = build_review_checklist_snapshot(lambda _item_id: True)
    record["review_checklist_complete"] = True
    assert "failed Governed" in validate_review_record(record, report_record=report_record)


def test_audit_paths_are_unique_for_rapid_successive_reports(monkeypatch):
    with TemporaryDirectory() as directory:
        monkeypatch.setattr("src.audit.AUDIT_DIR", Path(directory))
        review_record = {"approval_status": DRAFT_STATUS}
        payload = {
            "inputs": {"location": "Cairns"},
            "report_version": 1,
            "report_text": append_human_signoff("Draft", review_record),
            "human_review": review_record,
        }

        first = save_report_audit(payload)
        second = save_report_audit(payload)

        assert first != second
        assert len(list(Path(directory).glob("audit_*.json"))) == 2
