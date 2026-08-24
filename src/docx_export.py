import re
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.export_content import extract_report_metadata, plain_markdown_text
from src.markdown_tables import (
    is_markdown_table_row,
    is_markdown_table_separator,
    parse_markdown_table_row,
)

BODY_FONT = "Microsoft YaHei"
DRAFT_NOTICE = "DRAFT - NOT EMERGENCY ADVICE - HUMAN REVIEW REQUIRED"
DRAFT_NOTICE_DETAIL = (
    "Preparedness planning support only. This document does not provide live fire conditions, "
    "fire bans, evacuation orders or life-safety directions. The responsible organisation must "
    "review and approve it before formal use. Call 000 in a life-threatening emergency."
)


def _plain_markdown_text(text):
    return plain_markdown_text(text)


def _extract_meta_from_report(markdown_text):
    meta = extract_report_metadata(markdown_text)
    meta["generated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
    return meta


def _set_run_font(run, font_name=BODY_FONT):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_font(paragraph, size=10.5, color=None, bold=False):
    for run in paragraph.runs:
        _set_run_font(run)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)


def _set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    _set_run_font(run)
    run.font.size = Pt(9.5)
    run.bold = bold


def _add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def _configure_document(document):
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 22, (24, 33, 47)),
        ("Heading 1", 15, (127, 42, 25)),
        ("Heading 2", 12.5, (24, 33, 47)),
    ]:
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)


def _add_header_footer(document):
    section = document.sections[-1]
    header_para = section.header.paragraphs[0]
    header_para.text = "BushfireReadyGPT"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_font(header_para, size=8.5, color=(102, 112, 133))

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.add_run("Page ")
    _add_page_number(footer_para)
    _set_paragraph_font(footer_para, size=8.5, color=(102, 112, 133))


def _add_cover(document, meta):
    document.add_paragraph()
    kicker = document.add_paragraph("BushfireReadyGPT")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(kicker, size=11, color=(127, 42, 25), bold=True)

    title = document.add_paragraph(meta["title"])
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(title, size=22, color=(24, 33, 47), bold=True)

    for label, value in [
        ("Location", meta["location"]),
        ("Audience", meta["audience"]),
        ("Generated", meta["generated_at"]),
    ]:
        paragraph = document.add_paragraph(f"{label}: {value}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_font(paragraph, size=10.5, color=(52, 64, 84))

    document.add_paragraph()
    banner = document.add_paragraph(DRAFT_NOTICE)
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(banner, size=11, color=(180, 35, 24), bold=True)

    note = document.add_paragraph(DRAFT_NOTICE_DETAIL)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(note, size=9, color=(102, 112, 133))
    document.add_page_break()


def _add_metadata_table(document, meta):
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Location", meta["location"]),
        ("Audience", meta["audience"]),
        ("Generated", meta["generated_at"]),
    ]
    for row, (label, value) in zip(table.rows, rows):
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[1], value)
    document.add_paragraph()


def _add_markdown_table(document, table_lines):
    data = [
        cells
        for line in table_lines
        if not is_markdown_table_separator(line)
        if (cells := parse_markdown_table_row(line)) is not None
    ]
    if not data:
        return
    column_count = max(len(row) for row in data)
    if column_count > 4:
        _add_record_tables(document, data)
        return
    table = document.add_table(rows=len(data), cols=column_count)
    table.style = "Table Grid"
    for row_index, row_data in enumerate(data):
        for col_index in range(column_count):
            value = row_data[col_index] if col_index < len(row_data) else ""
            _set_cell_text(table.rows[row_index].cells[col_index], _plain_markdown_text(value), bold=row_index == 0)
        if row_index == 0:
            _set_row_flag(table.rows[row_index], "tblHeader")
        if sum(len(value) for value in row_data) <= 900:
            _set_row_flag(table.rows[row_index], "cantSplit")
    document.add_paragraph()


def _set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_row_flag(row, flag):
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn(f"w:{flag}")) is None:
        properties.append(OxmlElement(f"w:{flag}"))


def _add_record_tables(document, data):
    headers = data[0]
    for record_index, row_data in enumerate(data[1:], start=1):
        label = document.add_paragraph(f"Record {record_index}")
        label.paragraph_format.keep_with_next = True
        _set_paragraph_font(label, size=9.5, color=(127, 42, 25), bold=True)

        table = document.add_table(rows=len(headers), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for row_index, header in enumerate(headers):
            value = row_data[row_index] if row_index < len(row_data) else ""
            label_cell, value_cell = table.rows[row_index].cells
            label_cell.width = Cm(4.2)
            value_cell.width = Cm(12.0)
            _set_cell_shading(label_cell, "F6F8FB")
            _set_cell_text(label_cell, _plain_markdown_text(header), bold=True)
            _set_cell_text(value_cell, _plain_markdown_text(value))
            _set_row_flag(table.rows[row_index], "cantSplit")
            if row_index < len(headers) - 1:
                for cell in table.rows[row_index].cells:
                    cell.paragraphs[0].paragraph_format.keep_with_next = True
        document.add_paragraph()


def _add_markdown_body(document, markdown_text):
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue

        if is_markdown_table_row(line):
            table_lines = []
            while index < len(lines) and is_markdown_table_row(lines[index].strip()):
                table_lines.append(lines[index].strip())
                index += 1
            _add_markdown_table(document, table_lines)
            continue
        elif line.startswith("# "):
            heading = _plain_markdown_text(line)
            if heading.lower() != "bushfirereadygpt report":
                document.add_heading(heading, level=1)
        elif line.startswith("## "):
            heading = _plain_markdown_text(line)
            if heading.lower() == "human review sign-off":
                document.add_page_break()
            document.add_heading(heading, level=1)
        elif line.startswith("### "):
            document.add_heading(_plain_markdown_text(line), level=2)
        elif line.startswith(("- [ ]", "- [x]", "- [X]")):
            checked = "☑" if line[3].lower() == "x" else "☐"
            text = _plain_markdown_text(line[5:].strip())
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(f"{checked} {text}")
            _set_run_font(run)
        elif line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(_plain_markdown_text(line[2:]))
            _set_run_font(run)
        elif re.match(r"^\d+\.\s+", line):
            paragraph = document.add_paragraph(style="List Number")
            run = paragraph.add_run(_plain_markdown_text(re.sub(r"^\d+\.\s+", "", line)))
            _set_run_font(run)
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(_plain_markdown_text(line))
            _set_run_font(run)
        index += 1


def create_report_docx(markdown_text):
    meta = _extract_meta_from_report(markdown_text)
    document = Document()
    _configure_document(document)
    _add_cover(document, meta)

    document.add_section(WD_SECTION_START.NEW_PAGE)
    _add_header_footer(document)
    document.add_heading(meta["title"], level=1)
    _add_metadata_table(document, meta)
    _add_markdown_body(document, markdown_text)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
